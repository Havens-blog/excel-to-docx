#!/usr/bin/python

import tkinter as tk
from tkinter import filedialog
import os
from shutil import move, copy2
from datetime import datetime
from time import sleep
from subprocess import Popen
from functools import partial
from time import time
from docx import Document
from docx.enum.text import WD_BREAK
from docxtpl import DocxTemplate
from openpyxl import load_workbook


def column_to_name(colnum):
    result = ""
    while colnum > 0:
        colnum, remainder = divmod(colnum - 1, 26)
        result = chr(65 + remainder) + result
    return result


# Function to format Excel file
def format_excel(excel_path):
    excel_dir, excel_file = os.path.split(excel_path)
    excel_name, excel_ext = os.path.splitext(excel_file)

    new_excel_name = f"{excel_name}_gork{excel_ext}"
    new_excel_path = os.path.join(excel_dir, new_excel_name)

    # Use shutil.copy2 to ensure the same attributes (e.g., timestamps) are copied
    copy2(excel_path, new_excel_path)

    wb = load_workbook(new_excel_path, data_only=True)
    sheet = wb.active

    for i in range(sheet.max_row):
        for j in range(sheet.max_column):
            cell_type = sheet["%s%d" % (column_to_name(j + 1), i + 1)].number_format
            fin_type = ['_ * #,##0.00_ ;_ * \\-#,##0.00_ ;_ * "-"??_ ;_ @_ ']

            if str(cell_type) == fin_type[0]:
                if type(sheet["%s%d" % (column_to_name(j + 1), i + 1)].value) == str:
                    break
                else:
                    float_value = float(
                        sheet["%s%d" % (column_to_name(j + 1), i + 1)].value
                    )
                sheet["%s%d" % (column_to_name(j + 1), i + 1)] = f"{float_value:,.2f}"
    wb.save(new_excel_path)


def process_files():
    excel_path = file_entry_1.get()
    docx_template_path = file_entry_2.get()
    if excel_path and docx_template_path:
        result_label.config(text="Processing...", fg="blue")
        partial_process_func = partial(
            process_excel_to_docx, excel_path, docx_template_path, result_label
        )
        partial_process_func()
    else:
        result_label.config(
            text="Error: Please select both Excel and Word files.", fg="red"
        )


# Function to open the first file dialog (limit to xlsx files)
def open_file_dialog_1():
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    file_entry_1.delete(0, tk.END)
    file_entry_1.insert(0, file_path)


# Function to open the second file dialog (limit to docx files)
def open_file_dialog_2():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    file_entry_2.delete(0, tk.END)
    file_entry_2.insert(0, file_path)


# Function to process Excel to DOCX conversion
def process_excel_to_docx(excel_path, docx_template_path, result_label):
    start_time = time()  # Record the start time
    time_format = "{:.2f}s"

    excel_dir, excel_file = os.path.split(excel_path)
    excel_name, excel_ext = os.path.splitext(excel_file)

    new_excel_name = f"{excel_name}_gork{excel_ext}"
    new_excel_path = os.path.join(excel_dir, new_excel_name)

    current_time = datetime.now()
    folder_name = "{}-{}-{}-{}-{}-{}".format(
        str(current_time.year),
        str(current_time.month),
        str(current_time.day),
        str(current_time.hour),
        str(current_time.minute),
        str(current_time.second),
    )
    os.makedirs(folder_name)
    format_excel(excel_path)

    # Define the 'doc' variable here
    doc = DocxTemplate(docx_template_path)

    wb = load_workbook(new_excel_path, data_only=True, keep_vba=True)
    sheet = wb.active
    total_files = sheet.max_row - 1  # Subtract 1 for the header row

    processed_files = 0
    time_taken = 0
    rendered_docs_paths = []
    for row in sheet:
        sleep(0.01)
        rowMap = map(lambda x: x.value, row)
        if row[0].row == 1:
            title = list(rowMap)
        else:
            context = dict(zip(title, rowMap))
            ##处理时间数据
            time_data = context.get("time")
            if time_data:
                data_part = time_data.split(' ')[0]
                context["time"] = data_part
            ## todo 这里要指定文件名所在列
            filename = context["Id"]
      
            doc.render(context)
            doc.save("%s.docx" % filename)
            rendered_docs_paths.append("./%s/%s.docx" % (folder_name,filename))
            processed_files += 1
            time_taken = time() - start_time
            progress_message = f"Processing {filename} ({processed_files}/{total_files}) \n- Time taken: {time_format.format(time_taken)} \n- Lines processed: {processed_files}"
            result_label.config(text=progress_message, fg="blue")
            root.update()  # Update the GUI to show the progress
            move("%s.docx" % filename, "./%s" % folder_name)

    merged_document=Document()
    for i,path in enumerate(rendered_docs_paths):
        # 打开渲染后的文档
        sub_doc = Document(path)
        # 在第一个文档之前不添加分页符
        # if i > 0:
        #     merged_document.add_page_break()
        # 将渲染后的文档内容添加到合并文档中   
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
        for sum in range(0,3):
            merged_document.add_paragraph()
        # 保存合并后的文档
    #merged_document.save("./%s/%s" % (folder_name,'merged_document.docx'))

    os.remove(new_excel_path)
    completed_text = f"Processing completed. Check the output folder.\n- Time taken: {time_format.format(time_taken)} \n - Lines processed: {processed_files}\n - Average speed: {time_format.format(time_taken/processed_files)} per file"
    result_label.config(text=completed_text, fg="green")

    # Open the output directory after processing
    if os.path.exists(folder_name):
        Popen(["explorer", folder_name])  # Use "explorer" for Windows
    else:
        result_label.config(text="Output directory does not exist.", fg="red")


# Create the main Tkinter window
root = tk.Tk()
root.title("Excel to Word Converter")

# Create and configure GUI elements
excel_label = tk.Label(root, text="Select Excel File:")
excel_label.grid(row=0, column=0)

file_entry_1 = tk.Entry(root, width=40)
file_entry_1.grid(row=0, column=1)

file_button_1 = tk.Button(root, text="Browse", command=open_file_dialog_1)
file_button_1.grid(row=0, column=2)

docx_label = tk.Label(root, text="Select Word Template:")
docx_label.grid(row=1, column=0)

file_entry_2 = tk.Entry(root, width=40)
file_entry_2.grid(row=1, column=1)

file_button_2 = tk.Button(root, text="Browse", command=open_file_dialog_2)
file_button_2.grid(row=1, column=2)

process_button = tk.Button(root, text="Process", command=process_files)
process_button.grid(row=2, column=0, columnspan=3, pady=10)

result_label = tk.Label(root, text="", fg="black")
result_label.grid(row=3, column=0, columnspan=3)

hello_label = tk.Label(
    root,
    text="Developer: ChengYingHao\n Last Updated on: 2024",
    fg="blue",
)
hello_label.grid(row=4, column=0, columnspan=3)

# Center the buttons and panels in the window
root.update()
window_width = root.winfo_width()
window_height = root.winfo_height()
x_offset = (root.winfo_screenwidth() - window_width) // 2
y_offset = (root.winfo_screenheight() - window_height) // 2
root.geometry(f"{window_width}x{window_height}+{x_offset}+{y_offset}")

# Start the Tkinter main loop
root.mainloop()
